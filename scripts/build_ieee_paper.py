#!/usr/bin/env python3
"""
Build IEEE CONECCT-formatted revision of the medication recommendation paper.
Applies all review recommendations: structure, length, fixes, anonymization.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "final_paper-mehra_esfandiari_.docx"
OUT = ROOT / "final_paper-mehra_esfandiari_.docx"
BACKUP = ROOT / "final_paper-mehra_esfandiari_original.docx"

TITLE = (
    "Safe and Explainable Medication Recommendation via "
    "Transformer-GAT Fusion on MIMIC-III"
)

ABSTRACT = (
    "Medication recommendation from electronic health records must balance predictive "
    "accuracy with pharmacological safety and clinical interpretability. We present a "
    "Transformer- and graph-based framework that integrates longitudinal patient encoding, "
    "a dual-graph drug module (DDI-aware graph attention network and molecular substructure "
    "GNN), learnable gated fusion, and a severity-weighted multi-objective loss. DeepSHAP "
    "provides patient-specific explanations. Of seven specified architectural modules, five "
    "are implemented and validated on the MIMIC-III GAMENet cohort (5,430 patients, 153 "
    "medications); outcome monitoring and additional explainability methods are specified for "
    "future work. Across four random seeds, the model achieves Jaccard 0.2789±0.0085, "
    "F1-micro 0.4434±0.0083, PRAUC 0.4675±0.0054, and DDI rate 0.0000. Ablation confirms "
    "that the DDI penalty is the primary safety mechanism but reduces Jaccard from 0.3750 "
    "when disabled. A λ_DDI Pareto sweep shows a monotonic safety-accuracy trade-off. SHAP "
    "analysis reveals clinically coherent diagnosis-driven attributions. Results support the "
    "framework as a decision-support prototype prioritizing drug safety over raw accuracy."
)

KEYWORDS = (
    "Medication recommendation; Transformer; graph attention networks; "
    "drug-drug interaction; explainable AI; electronic health records"
)

REFERENCES = [
    "A. E. W. Johnson et al., \"MIMIC-III, a freely accessible critical care database,\" "
    "Sci. Data, vol. 3, p. 160035, 2016.",
    "X. Cai et al., \"Multiview clustering for EHR-based polypharmacy analysis: A review "
    "and benchmark,\" J. Biomed. Inform., vol. 131, p. 104113, 2022.",
    "E. Choi et al., \"RETAIN: An interpretable predictive model for healthcare using reverse "
    "time attention mechanism,\" in Proc. NeurIPS, 2016, pp. 3504–3512.",
    "Y. Zhang et al., \"LEAP: Learning to prescribe effective and safe treatment combinations "
    "for multimorbidity,\" in Proc. KDD, 2017, pp. 1315–1324.",
    "J. Shang et al., \"GAMENet: Graph augmented memory networks for recommending medication "
    "combinations,\" in Proc. AAAI, vol. 33, 2019, pp. 1126–1133.",
    "C. Yang et al., \"SafeDrug: Dual molecular graph encoders for safe drug recommendations,\" "
    "in Proc. IJCAI, 2021, pp. 3399–3405.",
    "R. Wu et al., \"Conditional generation net for medication recommendation,\" in Proc. WWW, "
    "2022, pp. 935–945.",
    "J. Shang et al., \"Pre-training of graph augmented transformers for medication "
    "recommendation (G-BERT),\" in Proc. IJCAI, 2019, pp. 5054–5060.",
    "Y. Li et al., \"BEHRT: Transformer for electronic health records,\" Sci. Rep., vol. 10, "
    "no. 1, p. 7155, 2020.",
    "Z. Zheng et al., \"MoleRec: Combinatorial drug recommendation with substructure "
    "representation learning,\" Bioinformatics, vol. 38, no. 17, pp. 4211–4219, 2022.",
    "L. Rasmy et al., \"Med-BERT: Pretrained contextualized embeddings on large-scale "
    "structured EHRs,\" NPJ Digit. Med., vol. 4, no. 1, p. 86, 2021.",
    "S. Jain and B. C. Wallace, \"Attention is not explanation,\" in Proc. NAACL-HLT, 2019, "
    "pp. 3543–3556.",
    "S. M. Lundberg and S.-I. Lee, \"A unified approach to interpreting model predictions,\" "
    "in Proc. NeurIPS, 2017, pp. 4765–4774.",
    "M. Sundararajan, A. Taly, and Q. Yan, \"Axiomatic attribution for deep networks,\" "
    "in Proc. ICML, 2017, pp. 3319–3328.",
    "Z. Ying et al., \"GNNExplainer: Generating explanations for graph neural networks,\" "
    "in Proc. NeurIPS, 2019, pp. 9240–9251.",
    "J. Futoma, J. Morris, and J. Lucas, \"A comparison of models for predicting early "
    "hospital readmissions,\" J. Biomed. Inform., vol. 56, pp. 229–238, 2015.",
    "Y. Zhang et al., \"Adverse drug event detection with deep neural networks,\" J. Biomed. "
    "Inform., vol. 106, p. 103432, 2020.",
    "D. S. Wishart et al., \"DrugBank 5.0: A major update for 2018,\" Nucleic Acids Res., "
    "vol. 46, no. D1, pp. D1074–D1082, 2018.",
    "A. Vaswani et al., \"Attention is all you need,\" in Proc. NeurIPS, 2017, pp. 5998–6008.",
    "S. Abnar and W. Zuidema, \"Quantifying attention flow in transformers,\" in Proc. ACL, "
    "2020, pp. 4190–4197.",
    "S. Bach et al., \"On pixel-wise explanations for non-linear classifier decisions by "
    "layer-wise relevance propagation,\" PLOS ONE, vol. 10, no. 7, p. e0130140, 2015.",
    "S. Bai, J. Z. Kolter, and V. Koltun, \"An empirical evaluation of generic convolutional "
    "and recurrent networks for sequence modeling,\" arXiv:1803.01271, 2018.",
    "Z. Obermeyer et al., \"Dissecting racial bias in an algorithm used to manage the health "
    "of populations,\" Science, vol. 366, no. 6464, pp. 447–453, 2019.",
    "J. Shang et al., \"MICRON: Medication change prediction using residual sequence "
    "learning,\" in Proc. IJCAI, 2021, pp. 3342–3348.",
    "P. Veličković et al., \"Graph attention networks,\" in Proc. ICLR, 2018.",
    "S. Bhoi, M. L. Li, and W. Hsu, \"PREMIER: Personalized recommendation for medical "
    "prescriptions from electronic records,\" arXiv:2008.13569, 2020.",
    "Y. Cai et al., \"Health recommender systems development, usage, and evaluation from 2010 "
    "to 2022: A scoping review,\" Int. J. Environ. Res. Public Health, vol. 19, no. 22, "
    "p. 15115, 2022.",
    "G. R. Langley et al., \"Towards a 21st-century roadmap for biomedical research and drug "
    "discovery,\" Drug Discov. Today, vol. 22, no. 2, pp. 327–339, 2017.",
    "M. Noshad, I. Jankovic, and J. H. Chen, \"Clinical recommender system: Predicting medical "
    "specialty diagnostic choices with neural network ensembles,\" arXiv:2007.12161, 2020.",
    "G. Liu et al., \"DNMDR: Dynamic networks and multi-view drug representations for safe "
    "medication recommendation,\" Eng. Appl. Artif. Intell., arXiv:2501.08572, 2025.",
    "K. Zhang et al., \"KEDRec-LM: A knowledge-distilled explainable drug recommendation large "
    "language model,\" arXiv:2502.20350, 2025.",
    "K. Singhal et al., \"Toward expert-level medical question answering with large language "
    "models,\" Nature Med., vol. 31, no. 3, pp. 943–950, 2025.",
    "Y. Ye et al., \"CT-PASMR: Personalized and safe medication recommendation based on "
    "convolutional neural network and transformer architecture,\" Eng. Appl. Artif. Intell., "
    "2025.",
]


def set_two_columns(section) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "720")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Twips(16838)  # A4
    section.page_width = Twips(11906)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    set_two_columns(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.bold = True
        h.font.size = Pt(10 if level > 1 else 11)


def add_para(doc, text, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
    return p


def add_table(doc, headers, rows, caption=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if caption:
        cap = doc.add_paragraph()
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.name = "Times New Roman"
    doc.add_paragraph()


def add_figure(doc, image_path: Path, caption: str, width=Inches(3.2)) -> None:
    if not image_path.exists():
        add_para(doc, f"[Insert {image_path.name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.name = "Times New Roman"
    doc.add_paragraph()


def extract_figures(backup: Path, dest_dir: Path) -> dict[str, Path]:
    dest_dir.mkdir(parents=True, exist_ok=True)
    mapping = {}
    if not backup.exists():
        return mapping
    import zipfile

    with zipfile.ZipFile(backup) as z:
        for name in z.namelist():
            if name.startswith("word/media/image") and name.endswith(".png"):
                fname = Path(name).name
                out = dest_dir / fname
                out.write_bytes(z.read(name))
                mapping[fname] = out
    return mapping


def build_document(figures: dict[str, Path] | None = None) -> Document:
    doc = Document()
    style_doc(doc)

    # Title (anonymized for double-blind review)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    add_para(doc, "Abstract—" + ABSTRACT, italic=False, space_after=4)
    add_para(doc, "Index Terms—" + KEYWORDS, italic=True, space_after=8)

    # I. INTRODUCTION
    add_heading(doc, "I. INTRODUCTION", level=1)
    intro = [
        "The widespread adoption of electronic health records (EHRs) has created opportunities "
        "for intelligent clinical decision support in medication recommendation [1]. Translating "
        "longitudinal EHR data into safe and interpretable drug sets remains challenging because "
        "of comorbidity complexity, temporal disease dynamics, and drug-drug interaction (DDI) "
        "risk in polypharmacy populations [2].",
        "Sequential models such as RETAIN [3] and LEAP [4] established temporal EHR modeling. "
        "Graph-augmented methods including GAMENet [5], SafeDrug [6], and COGNet [7] integrated "
        "pharmacological knowledge. Transformer-based approaches including G-BERT [8] and BEHRT [9] "
        "demonstrated long-range dependency modeling. Recent work extends molecular and dynamic "
        "graph representations [10], [26], [30], [33]. Yet no existing system jointly addresses "
        "longitudinal modeling, molecular safety, faithful explainability, and unified multi-objective "
        "optimization within one trainable framework.",
        "We propose a Transformer-GAT architecture with learnable gated fusion, a dual-graph drug "
        "module, severity-weighted DDI loss, and DeepSHAP explainability. Seven modules are specified "
        "mathematically; five are implemented and experimentally validated on MIMIC-III. Outcome "
        "monitoring and additional explainability methods (Integrated Gradients, LRP, GNNExplainer) "
        "are architectural specifications for future work.",
        "Contributions: (1) a validated Transformer-GAT framework with gated fusion and dual-graph "
        "drug encoding; (2) a four-term multi-objective loss with active L_rec and L_DDI terms; "
        "(3) patient-specific DeepSHAP explanations; (4) four-seed MIMIC-III evaluation with ablation, "
        "Pareto analysis, and transparent safety-accuracy reporting.",
    ]
    for para in intro:
        add_para(doc, para)

    # II. RELATED WORK
    add_heading(doc, "II. RELATED WORK", level=1)
    rw = [
        "LEAP [4] and GAMENet [5] model visit sequences with LSTM/GRU and DDI graphs respectively. "
        "SafeDrug [6] and MoleRec [10] encode molecular substructures but offer limited explainability. "
        "COGNet [7] models medication continuity via copy-and-generate decoding. Transformer models "
        "BEHRT [9], Med-BERT [11], and G-BERT [8] improve representation learning; CT-PASMR [33] "
        "combines CNN-Transformer encoding with graph attention for safety-aware recommendation.",
        "Explainability remains a gap: attention weights are not faithful explanations [12], whereas "
        "SHAP [13], Integrated Gradients [14], and GNNExplainer [15] provide stronger attribution "
        "guarantees. Outcome prediction (readmission [16], adverse events [17]) is typically studied "
        "separately from prescription generation.",
        "Table I summarizes capability coverage. The proposed method is the only entry combining "
        "Transformer temporal encoding, molecular graphs, severity-weighted DDI modeling, and gradient-based "
        "explainability; outcome monitoring and fairness evaluation are specified but not experimentally "
        "validated in this study.",
    ]
    for para in rw:
        add_para(doc, para)

    add_table(
        doc,
        ["Model", "Temporal", "DDI", "Molecular", "Explain.", "Outcome", "Fairness"],
        [
            ["LEAP", "LSTM", "✗", "✗", "✗", "✗", "✗"],
            ["GAMENet", "GRU", "~", "✗", "✗", "✗", "✗"],
            ["SafeDrug", "GRU", "✓", "✓", "✗", "✗", "✗"],
            ["MoleRec", "GRU", "✓", "✓", "✗", "✗", "✗"],
            ["G-BERT", "Trans.", "~", "~", "Attn", "✗", "✗"],
            ["Proposed*", "Trans.", "✓", "✓", "SHAP†", "~‡", "✗"],
        ],
        "TABLE I. FEATURE COMPARISON. ✓ FULLY ADDRESSED; ~ PARTIAL; ✗ ABSENT. "
        "*FIVE OF SEVEN MODULES VALIDATED. †DEEPSHAP ONLY. ‡OUTCOME MODULE SPECIFIED, NOT EVALUATED.",
    )

    # III. PROPOSED METHOD
    add_heading(doc, "III. PROPOSED METHOD", level=1)
    method = [
        "Each patient p_i has T visits with diagnosis, medication, and lab multi-hot vectors. "
        "Visit embedding e_t = GELU(W_a·[Emb_d⊕Emb_m⊕Emb_l]+b_a) is augmented with sinusoidal "
        "positional encoding and processed by a 4-layer, 8-head Transformer (d_model=256) [19] "
        "to obtain patient representation h_patient.",
        "Dual-graph drug encoding combines a 3-layer DDI-aware GAT over DrugBank 5.0 [18] with "
        "severity weights w_uv∈{0.1,0.5,1.0} and a molecular substructure GNN on SMILES-derived "
        "graphs (RDKit). Drug embeddings are fused as d_i = h_i^(GAT)⊕m_i. Learnable gated fusion "
        "integrates patient and drug views: g=σ(W_g[h_patient‖h_drug]+b_g), "
        "h_fused=g⊙h_patient+(1−g)⊙h_drug.",
        "The recommendation head outputs ŷ∈[0,1]^|M| via sigmoid. The unified loss is "
        "L_total=λ₁L_rec+λ₂L_DDI+λ₃L_outcome+λ₄L_xai. In experiments, L_rec is multi-label BCE "
        "(with per-drug pos_weight) and L_DDI=Σ w_uv·ŷ_u·ŷ_v penalizes co-prescribing interacting "
        "drugs. L_outcome and L_xai are defined architecturally but evaluate to zero because the "
        "TCN outcome head and Integrated Gradients are not yet implemented.",
        "DeepSHAP [13] provides patient-specific attributions computed against each patient's "
        "top-predicted medications (not pooled across all drugs). Additional explainability methods "
        "(attention rollout [20], LRP [21], GNNExplainer [15], counterfactuals) and the TCN outcome "
        "module [22] are specified for extension.",
    ]
    for para in method:
        add_para(doc, para)

    figs = figures or {}
    add_figure(
        doc,
        figs.get("image1.png", Path()),
        "Fig. 1. System architecture overview (seven-module framework).",
    )

    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Transformer layers / heads", "4 / 8"],
            ["d_model", "256"],
            ["GAT layers", "3"],
            ["Dropout", "0.30"],
            ["Learning rate (AdamW)", "1×10⁻⁴"],
            ["Batch size", "64"],
            ["λ₁, λ₂, λ₃, λ₄", "1.0, 0.5, 0.3, 0.1"],
        ],
        "TABLE II. KEY HYPERPARAMETERS.",
    )

    # IV. EXPERIMENTS AND RESULTS
    add_heading(doc, "IV. EXPERIMENTS AND RESULTS", level=1)
    exp = [
        "We evaluate on the MIMIC-III GAMENet cohort: 5,430 patients (3,800/814/816 train/val/test), "
        "1,958 diagnosis codes, 153 ATC-4 medications [1]. Preprocessing maps ICD-9 diagnoses, NDC "
        "drugs, and labs to visit-level multi-hot vectors; DDI edges come from DrugBank 5.0 [18]. "
        "Patient-level splitting prevents leakage. Results are mean±std over seeds {42,1,123,2024}.",
        "Baselines include LEAP, RETAIN, GAMENet, SafeDrug, MoleRec (with PubChem SMILES graphs, "
        "82.8% ATC-4 coverage), MICRON [24], G-BERT, and COGNet. Implementation uses PyTorch 2.0, "
        "PyTorch Geometric, SHAP 0.42, and AdamW with gradient clipping (norm 1.0).",
        "Table III reports main results. The proposed model achieves DDI rate 0.0000 across all seeds "
        "but lower Jaccard than sequence baselines (best baseline GAMENet: 0.3826). PRAUC (0.4675) "
        "exceeds baselines, indicating better ranking despite lower set overlap at threshold 0.5.",
    ]
    for para in exp:
        add_para(doc, para)

    add_figure(
        doc,
        figs.get("image4.png", Path()),
        "Fig. 2. Comparative performance across baseline models.",
    )
    add_figure(
        doc,
        figs.get("image5.png", Path()),
        "Fig. 3. DDI rate comparison across models.",
    )

    add_table(
        doc,
        ["Model", "Jaccard↑", "F1-micro↑", "PRAUC↑", "DDI↓"],
        [
            ["RETAIN", "0.3552", "0.5197", "0.1964", "0.0526"],
            ["LEAP", "0.3601", "0.5268", "0.1792", "0.0599"],
            ["GAMENet", "0.3826", "0.5570", "0.2146", "0.0577"],
            ["MoleRec", "0.2498", "0.3761", "—", "0.0009"],
            ["Proposed", "0.2789", "0.4434", "0.4675", "0.0000"],
        ],
        "TABLE III. MAIN RESULTS ON MIMIC-III (PROPOSED: MEAN±STD OVER 4 SEEDS).",
    )

    results2 = [
        "Ablation (Table IV): removing L_DDI raises Jaccard to 0.3750 but introduces DDI rate 0.0397. "
        "Removing the Transformer slightly improves Jaccard (0.2975 vs. 0.2738), consistent with "
        "short visit sequences (T≤5) and fixed training budget; diagnostic loss curves show no "
        "overfitting (train-val gap <0.003).",
        "Safety-adjusted effectiveness (SAE=Jaccard×(1−DDI)) narrows but does not close the gap "
        "versus baselines at λ₂=0.5. A λ_DDI sweep (Table V) shows monotonic trade-off: λ=0 yields "
        "best Jaccard (0.3750, DDI 0.0392); λ≥0.3 achieves zero DDI with Jaccard≈0.27–0.28.",
        "DeepSHAP on 100 test patients shows diagnosis features dominate (e.g., long-term anticoagulant "
        "use V58.61, severe sepsis 995.92), consistent with ICU prescribing. 95.6% of patients have "
        "unique top-10 drug sets, indicating genuine personalization.",
    ]
    for para in results2:
        add_para(doc, para)

    add_figure(
        doc,
        figs.get("image9.png", Path()),
        "Fig. 4. Top-20 SHAP feature importance (diagnosis codes dominate).",
    )

    add_table(
        doc,
        ["Variant", "Jaccard", "F1-micro", "DDI"],
        [
            ["Full model", "0.2738", "0.4355", "0.0000"],
            ["w/o Transformer", "0.2975", "0.4768", "0.0000"],
            ["w/o GAT", "0.2589", "0.4207", "0.0000"],
            ["w/o DDI loss", "0.3750", "0.5538", "0.0397"],
        ],
        "TABLE IV. ABLATION (SEED 42).",
    )

    add_table(
        doc,
        ["λ_DDI", "Jaccard", "DDI", "SAE"],
        [
            ["0.0", "0.3750", "0.0392", "0.3603"],
            ["0.1", "0.3017", "0.0003", "0.3016"],
            ["0.5*", "0.2738", "0.0000", "0.2738"],
            ["1.0", "0.2706", "0.0000", "0.2706"],
        ],
        "TABLE V. PARETO SWEEP. *DEFAULT OPERATING POINT.",
    )

    # V. DISCUSSION
    add_heading(doc, "V. DISCUSSION", level=1)
    disc = [
        "The proposed framework prioritizes pharmacological safety: zero predicted DDIs are achieved "
        "at a measurable accuracy cost versus GAMENet and LEAP. This is a deliberate decision-support "
        "trade-off, not evidence of superior overall recommendation quality. Practitioners should select "
        "λ_DDI based on institutional risk tolerance using the Pareto sweep.",
        "Gated fusion enables per-patient weighting of clinical trajectory versus drug-graph signals—"
        "a mechanism absent from fixed concatenation fusion. The w/o Transformer result suggests that "
        "at this cohort scale, simpler encoders may converge faster under a fixed epoch budget [9].",
        "Limitations: (1) single-center MIMIC-III ICU cohort; (2) only DeepSHAP validated among seven "
        "specified explainability methods; (3) outcome module not evaluated; (4) no statistical significance "
        "tests versus baselines; (5) two high-frequency drugs show near-zero recall under strong DDI "
        "regularization, suggesting need for curriculum or confidence-aware penalty weighting; "
        "(6) fairness not empirically assessed despite design considerations [23].",
        "Ethics: all data are de-identified. Outputs require physician review and are not autonomous "
        "prescriptions. Federated training and regulatory compliance remain future engineering tasks.",
    ]
    for para in disc:
        add_para(doc, para)

    # VI. CONCLUSION
    add_heading(doc, "VI. CONCLUSION", level=1)
    concl = [
        "We presented a Transformer-GAT medication recommendation framework with dual-graph drug "
        "encoding, gated fusion, DDI-aware multi-objective training, and DeepSHAP explainability. "
        "On MIMIC-III, the validated implementation achieves DDI rate 0.0000 (four seeds) with "
        "Jaccard 0.2789±0.0085. Ablation and Pareto analysis characterize the safety-accuracy "
        "frontier transparently. Future work includes full module implementation, MIMIC-IV/eICU "
        "validation, significance testing, and prospective clinical evaluation.",
    ]
    for para in concl:
        add_para(doc, para)

    # REFERENCES
    add_heading(doc, "REFERENCES", level=1)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        run = p.add_run(f"[{i}] {ref}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)

    return doc


def main():
    if SRC.exists() and not BACKUP.exists():
        shutil.copy2(SRC, BACKUP)
        print(f"Backup saved: {BACKUP}")

    fig_dir = ROOT / ".paper_figures"
    figures = extract_figures(BACKUP, fig_dir)

    doc = build_document(figures)
    doc.save(OUT)
    print(f"Revised paper written: {OUT}")

    # stats
    text = "\n".join(p.text for p in doc.paragraphs)
    words = len(text.split())
    abs_words = len(ABSTRACT.split())
    print(f"Approximate word count: {words}")
    print(f"Abstract word count: {abs_words}")
    print(f"Reference count: {len(REFERENCES)}")


if __name__ == "__main__":
    main()
